import gtfs_kit as gk
import pandas as pd
import streamlit as st
import os
import re
from utils import group_dates_by_timetables, organize_times_by_hour
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import zipfile
import tempfile
import copy

st.set_page_config(page_title="Stop Passenger Information Generator", page_icon=":busstop:")

def load_feed(file):
    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(file.getvalue())
        tmp_path = tmp.name
    feed = gk.read_feed(tmp_path, dist_units='km').clean()
    return feed

def get_first_and_subsequent_weeks(feed):
    first_week = gk.calendar.get_first_week(feed, as_date_obj=True)
    valid_mondays = []
    if first_week:
        valid_mondays.append(first_week[0])
        k = 2
        while True:
            week = gk.calendar.get_week(feed, k, as_date_obj=True)
            if not week:
                break
            valid_mondays.append(week[0])
            k += 1
    return valid_mondays

def get_route_timetable(feed, route_id, week_dates, columns=None):
    timetable = gk.routes.build_route_timetable(feed, route_id, week_dates)
    if columns:
        timetable = timetable[columns]
    return timetable

def check_stop_sequences(trips, stop_times):
    trip_stop_sequences = stop_times.groupby('trip_id')['stop_id'].apply(tuple).reset_index()
    trips = trips.merge(trip_stop_sequences, on='trip_id')
    headsign_stop_sequences = trips.groupby('trip_headsign')['stop_id'].unique().reset_index()
    headsign_stop_sequences['stop_sequence_str'] = headsign_stop_sequences['stop_id'].apply(lambda x: ' -> '.join(map(str, x)))
    return headsign_stop_sequences

def classify_stops(timetable):
    timetable = timetable.sort_values(['trip_id', 'date', 'stop_sequence'])
    timetable['stop_type'] = 'Stop'
    timetable.loc[timetable.groupby(['trip_id', 'date'])['stop_sequence'].idxmin(), 'stop_type'] = 'Start'
    timetable.loc[timetable.groupby(['trip_id', 'date'])['stop_sequence'].idxmax(), 'stop_type'] = 'Finish'
    return timetable

def index_headsigns(timetable):
    headsign_index_dict = {}
    headsign_counts = timetable.groupby(['stop_id', 'trip_headsign']).size().reset_index(name='count')
    headsign_counts['headsign_index'] = headsign_counts.groupby('stop_id')['count'].rank("dense", ascending=False).astype(int)
    headsign_index_dict = headsign_counts.groupby('stop_id').apply(
        lambda x: dict(zip(x['headsign_index'], x['trip_headsign']))
    ).to_dict()
    headsign_index_map = headsign_counts.set_index(['stop_id', 'trip_headsign'])['headsign_index']
    timetable['headsign_index'] = timetable.set_index(['stop_id', 'trip_headsign']).index.map(headsign_index_map)
    return headsign_index_dict

def sanitize_filename(filename):
    return re.sub(r'[\\/*?:"<>|]', "_", filename)

def set_cell_background(cell, color):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._element.tcPr
    tcBorders = OxmlElement('w:tcBorders')
    
    def add_border(border_name, border_properties):
        border = OxmlElement(f'w:{border_name}')
        for key in border_properties:
            border.set(qn(f'w:{key}'), str(border_properties[key]))
        tcBorders.append(border)
    
    if top:
        add_border('top', top)
    if bottom:
        add_border('bottom', bottom)
    if left:
        add_border('left', left)
    if right:
        add_border('right', right)

    tc.append(tcBorders)

def adjust_font_size(cell, num_columns, bold=False):
    base_font_size = 12 - (num_columns // 5)
    if bold:
        base_font_size -= 1
    font_size = max(base_font_size, 6)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            if bold:
                run.font.bold = True

def create_styled_table(dataframe, headsign_colors=None, headsign_index_map=None):
    doc = Document()
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    num_columns = len(dataframe.columns)

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(dataframe.columns):
        hdr_cells[i].text = col_name
        adjust_font_size(hdr_cells[i], num_columns, bold=True)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[i].paragraphs[0].paragraph_format.keep_with_next = True
        set_cell_background(hdr_cells[i], 'E3EAEE')
        set_cell_border(hdr_cells[i],
                        top={"sz": "12", "val": "single", "color": "000000", "space": "0"},
                        bottom={"sz": "12", "val": "single", "color": "000000", "space": "0"},
                        left={"sz": "12", "val": "single", "color": "000000", "space": "0"} if i == 0 else None,
                        right={"sz": "12", "val": "single", "color": "000000", "space": "0"} if i == num_columns - 1 else None)

    num_rows = len(dataframe.values)
    for row_index, row_data in enumerate(dataframe.values):
        row_cells = table.add_row().cells
        for col_index, val in enumerate(row_data):
            cell = row_cells[col_index]
            cell.text = str(val)
            adjust_font_size(cell, num_columns)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            if headsign_index_map is not None and headsign_colors is not None:
                headsign_index = headsign_index_map.iloc[row_index, col_index]
                if headsign_index in headsign_colors:
                    set_cell_background(cell, headsign_colors[headsign_index])

            set_cell_border(cell,
                            top={"sz": "12", "val": "single", "color": "000000", "space": "0"} if row_index == 0 else {"sz": "0", "val": "none", "color": "000000", "space": "0"},
                            bottom={"sz": "12", "val": "single", "color": "000000", "space": "0"} if row_index == num_rows - 1 else {"sz": "0", "val": "none", "color": "000000", "space": "0"},
                            left={"sz": "12", "val": "single", "color": "000000", "space": "0"} if col_index == 0 else {"sz": "6", "val": "single", "color": "000000", "space": "0"},
                            right={"sz": "12", "val": "single", "color": "000000", "space": "0"} if col_index == num_columns - 1 else {"sz": "6", "val": "single", "color": "000000", "space": "0"})

        if row_index < num_rows - 1:
            for cell in row_cells:
                cell.paragraphs[0].paragraph_format.keep_with_next = True

    return table

def create_legend_table(headsign_colors, headsign_index_dict):
    table = Document().add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for key in sorted(headsign_index_dict.keys()):
        if key == 1:
            continue
        if key in headsign_colors:
            row_cells = table.add_row().cells
            set_cell_background(row_cells[0], headsign_colors[key])
            row_cells[0].width = Cm(0.75)
            set_cell_border(row_cells[0], 
                            top={"sz": "6", "val": "single", "color": "000000", "space": "0"},
                            bottom={"sz": "6", "val": "single", "color": "000000", "space": "0"},
                            left={"sz": "6", "val": "single", "color": "000000", "space": "0"},
                            right={"sz": "6", "val": "single", "color": "000000", "space": "0"})
            row_cells[1].text = str(headsign_index_dict[key])
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    return table

def insert_elements_at_placeholders(doc, elements, placeholder):
    placeholder_locs = []

    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            placeholder_locs.append(paragraph)
    
    for section in doc.sections:
        for header_footer in [section.header, section.footer]:
            for paragraph in header_footer.paragraphs:
                if placeholder in paragraph.text:
                    placeholder_locs.append(paragraph)

    for placeholder_paragraph in placeholder_locs:
        parent = placeholder_paragraph._element.getparent()
        idx = parent.index(placeholder_paragraph._element)
        placeholder_paragraph.clear()
        for element in elements:
            parent.insert(idx, copy.deepcopy(element._element))
            idx += 1

def generate_word_documents(timetable, route_id, feed, headsign_index_dict, status_container, template_file):
    doc_files = []
    route_short_name = feed.routes[feed.routes['route_id'] == route_id]['route_short_name'].values[0]
    headsign_colors = {
        1: 'FFFFFC',
        2: 'BED5FF',
        3: 'FEDAAA',
        4: 'BFFBB7',
        5: 'FFC6FF',
        6: 'FFFE88',
        7: 'D9D5FF',
        8: 'FFC8C8',
        9: 'A3F6FE'
    }

    num_stops = len(timetable['stop_id'].unique())

    for i, stop_id in enumerate(timetable['stop_id'].unique()):
        grouped_timetables = group_dates_by_timetables(timetable, stop_id)
        stop_name = feed.stops[feed.stops['stop_id'] == stop_id]['stop_name'].values[0]
        first_headsign_index = timetable[timetable['stop_id'] == stop_id]['headsign_index'].min()
        direction = timetable[(timetable['stop_id'] == stop_id) & (timetable['headsign_index'] == first_headsign_index)]['trip_headsign'].values[0]

        doc_filename = sanitize_filename(f"{route_id}_{stop_id}.docx")
        doc = DocxTemplate(template_file)

        context = {
            'route': route_short_name,
            'stop': stop_name,
            'direction': direction
        }
        doc.render(context)

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)

        rendered_doc = Document(doc_buffer)

        elements = []
        for group_name, (dates, group_timetable) in grouped_timetables.items():
            group_title = create_title(group_name)
            organized_timetable, organized_headsign_index = organize_times_by_hour(group_timetable['arrival_time'], group_timetable['headsign_index'])
            table = create_styled_table(organized_timetable, headsign_colors, organized_headsign_index)
            elements.append(group_title)
            elements.append(table)

        insert_elements_at_placeholders(rendered_doc, elements, '[TABLE_PLACEHOLDER]')
        legend_table = create_legend_table(headsign_colors, headsign_index_dict[stop_id])
        insert_elements_at_placeholders(rendered_doc, [legend_table], '[LEGEND_PLACEHOLDER]')

        output_buffer = io.BytesIO()
        rendered_doc.save(output_buffer)
        output_buffer.seek(0)
        doc_files.append((doc_filename, output_buffer))

        status_container.update(label=f"Processed {i + 1} of {num_stops} stops", state="running")

    status_container.update(label="All documents have been processed successfully!", state="complete")
    
    return doc_files

def create_title(text):
    doc = Document()
    title = doc.add_paragraph(text)
    title.style = 'Heading 2'
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.keep_with_next = True
    return title

def create_zip_file(doc_files):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w') as z:
        for doc_file, buffer in doc_files:
            z.writestr(doc_file, buffer.getvalue())
    return zip_buffer

st.title("GTFS Route Timetable Generator")

gtfs_file = st.file_uploader("Upload GTFS zip file", type="zip")
template_file = st.file_uploader("Upload Word template", type="docx")

if gtfs_file and template_file:
    feed = load_feed(gtfs_file)
    valid_mondays = get_first_and_subsequent_weeks(feed)

    if valid_mondays:
        selected_monday = st.selectbox("Select Start Date (Monday):", options=valid_mondays)
        route_options = feed.routes[['route_id', 'route_short_name', 'route_long_name', 'route_color', 'route_text_color']].drop_duplicates()
        route_options['route_label'] = route_options['route_short_name'] + " - " + route_options['route_long_name']
        route_dict = dict(zip(route_options['route_label'], route_options['route_id']))
        selected_route = st.selectbox("Select Route:", options=route_dict.keys())
        selected_route_id = route_dict[selected_route]
        route_info = route_options[route_options['route_id'] == selected_route_id].iloc[0]
        route_color = route_info['route_color']
        st.write(f"Route Color: {route_color}")

        additional_routes = []
        combine_lines = st.checkbox("Combine with other lines")
        if combine_lines:
            route_ids_to_combine = st.multiselect("Select additional lines to combine", options=route_options['route_id'].tolist())
            for additional_route_id in route_ids_to_combine:
                alias = st.text_input(f"Enter alias for {additional_route_id}")
                additional_routes.append((alias, additional_route_id))

        if st.button("Get Timetable"):
            route_id = selected_route_id
            week_dates = [(selected_monday + pd.DateOffset(days=i)).strftime('%Y%m%d') for i in range(7)]
            
            with st.status("Fetching timetable data...") as status_container:
                timetable = get_route_timetable(feed, route_id, week_dates, ["route_id", "trip_id", "trip_headsign", "arrival_time", "stop_id", "stop_sequence", "date"])

                if timetable.empty:
                    status_container.update(label="No data available for the selected route and date.", state="error")
                else:
                    if additional_routes:
                        for alias, additional_route_id in additional_routes:
                            additional_timetable = get_route_timetable(feed, additional_route_id, week_dates, ["route_id", "trip_id", "trip_headsign", "arrival_time", "stop_id", "stop_sequence", "date"])
                            additional_timetable['trip_headsign'] = alias + ' - ' + additional_timetable['trip_headsign']
                            timetable = pd.concat([timetable, additional_timetable])

                    timetable = classify_stops(timetable)
                    timetable = timetable[timetable['stop_type'] != 'Finish']

                    timetable = timetable.merge(feed.stops[['stop_id', 'stop_name']], on='stop_id', how='left')
                    timetable['stop_full'] = timetable['stop_id'] + " - " + timetable['stop_name']

                    headsign_index_dict = index_headsigns(timetable)

                    status_container.update(label="Generating Word documents...", state="running")
                    
                    doc_files = generate_word_documents(timetable, route_id, feed, headsign_index_dict, status_container, template_file)
                    zip_buffer = create_zip_file(doc_files)

                    st.download_button(
                        label="Download all Word documents as a zip file",
                        data=zip_buffer.getvalue(),
                        file_name="timetables.zip",
                        mime="application/zip"
                    )
    else:
        st.write("No valid Mondays found within the feed's date range.")
